# -*- coding: utf-8 -*-
"""
Python Native Messaging Host (암호감지 강화 + OLE + PDF 보강 + 엄격 모드)
- stdin/stdout (4바이트 LE 길이 + JSON) 프로토콜
- 텍스트/파일 PII 탐지(정규식 + Luhn/RRN/BizReg)
- DOCX/XLSX/PDF 암호 파일 → files[].error="password_protected"
- 레거시 OLE(.xls/.doc) 암호 파일도 password_protected로 표기
- PDF: /Encrypt 바이트 스캔 + (선택) PyPDF2.is_encrypted 확인
- 평문은 반환/저장하지 않음. (옵션) 해시만

필요 패키지(권장):
  pip install pdfminer.six python-docx openpyxl regex chardet msoffcrypto-tool olefile PyPDF2
"""
import sys, json, struct, base64, hashlib, io, re, math, zipfile
from typing import Dict, List, Any, Tuple

VERSION = "1.3.1-ooxml+ole+pdf-strict"

# ===== 옵션: OOXML 암호 감지 엄격 모드(권장 True) =====
STRICT_OOXML_PASSWORD_HEURISTIC = True

# -------- 선택 모듈 --------
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    from openpyxl import load_workbook
    from openpyxl.utils.exceptions import InvalidFileException
except Exception:
    load_workbook = None
    class InvalidFileException(Exception): ...
try:
    import regex as safe_re  # catastrophic backtracking 방지
except Exception:
    safe_re = re
try:
    import chardet
except Exception:
    chardet = None
# OLE 암호 판별용
try:
    import msoffcrypto
except Exception:
    msoffcrypto = None
try:
    import olefile  # msoffcrypto 내부에서 사용
except Exception:
    olefile = None
# PDF 보강용 (선택)
try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None


# ==========================
# 커스텀 예외
# ==========================
class PasswordProtectedError(Exception):
    """파일이 암호로 보호되어 텍스트 추출 불가"""
    pass


# ==========================
# Native Messaging I/O
# ==========================
def _read_msg() -> Dict[str, Any]:
    raw_len = sys.stdin.buffer.read(4)
    if not raw_len:
        sys.exit(0)
    msg_len = struct.unpack('<I', raw_len)[0]
    if msg_len <= 0 or msg_len > (256 * 1024 * 1024):
        sys.exit(0)
    data = sys.stdin.buffer.read(msg_len)
    return json.loads(data)

def _send_msg(obj: Dict[str, Any]):
    data = json.dumps(obj, ensure_ascii=False).encode('utf-8')
    sys.stdout.buffer.write(struct.pack('<I', len(data)))
    sys.stdout.buffer.write(data)
    sys.stdout.buffer.flush()


# ==========================
# 유틸/검증
# ==========================
def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()

def sha256_text(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8", "ignore")).hexdigest()

def luhn_check(number: str) -> bool:
    digits = [int(d) for d in re.sub(r'\D', '', number)]
    if len(digits) < 13:
        return False
    checksum = 0
    odd = True
    for d in reversed(digits):
        if odd:
            checksum += d
        else:
            dd = d * 2
            checksum += dd - 9 if dd > 9 else dd
        odd = not odd
    return (checksum % 10) == 0

def shannon_entropy(s: str) -> float:
    if not s:
        return 0.0
    from collections import Counter
    c = Counter(s); n = len(s)
    return -sum((cnt/n) * math.log2(cnt/n) for cnt in c.values())

def maybe_decode_text(b: bytes) -> str:
    if not b:
        return ""
    enc = "utf-8"
    if chardet:
        try:
            enc = chardet.detect(b).get("encoding") or "utf-8"
        except Exception:
            pass
    try:
        return b.decode(enc, errors="replace")
    except Exception:
        return b.decode("utf-8", errors="replace")

def rrn_valid(rrn: str) -> bool:
    d = re.sub(r"\D", "", rrn)
    if len(d) != 13:
        return False
    mm = int(d[2:4]); dd = int(d[4:6])
    if not (1 <= mm <= 12 and 1 <= dd <= 31):
        return False
    w = [2,3,4,5,6,7,8,9,2,3,4,5]
    s = sum(int(d[i]) * w[i] for i in range(12))
    check = (11 - (s % 11)) % 10
    return check == int(d[12])

def bizreg_valid(biz: str) -> bool:
    d = re.sub(r"\D", "", biz)
    if len(d) != 10:
        return False
    w = [1,3,7,1,3,7,1,3,5]
    s = sum(int(d[i]) * w[i] for i in range(9))
    s += (int(d[8]) * 5) // 10
    check = (10 - (s % 10)) % 10
    return check == int(d[9])


# ==========================
# 정규식 룰팩 (기본판)
# ==========================
REGEX = {
    "phone":  safe_re.compile(r"\b01[016789]-?\d{3,4}-?\d{4}\b"),
    "email":  safe_re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", safe_re.I),

    # 기본판: 하이픈형 + 13자리 연속 모두 후보(후속 버전에서 보수화 가능)
    "rrn":    safe_re.compile(r"\b\d{6}-\d{7}\b|\b\d{13}\b"),
    "bizreg": safe_re.compile(r"\b\d{3}-?\d{2}-?\d{5}\b"),

    "card_raw": safe_re.compile(r"(?:\d[ -]?){13,19}"),
    "account":  safe_re.compile(r"\b\d{8,14}\b"),

    "ipv4":   safe_re.compile(r"\b(?:(?:25[0-5]|2[0-4]\d|1?\d?\d)\.){3}(?:25[0-5]|2[0-4]\d|1?\d?\d)\b"),
    "ipv6":   safe_re.compile(r"\b(?:[A-F0-9]{1,4}:){7}[A-F0-9]{1,4}\b", safe_re.I),
    "zip5":   safe_re.compile(r"\b\d{5}\b"),
    "dob":    safe_re.compile(r"\b(19|20)\d{2}\s*(?:[.\-/\uFF0E]|년)\s*(?:0?[1-9]|1[0-2])\s*(?:[.\-/\uFF0E]|월)\s*(?:0?[1-9]|[12]\d|3[01])\s*(?:일)?\b"),

    "secret_prefix": safe_re.compile(r"\b(?:sk-[A-Za-z0-9]{16,}|AKIA[0-9A-Z]{16}|xoxb-[0-9A-Za-z-]{20,})\b"),
}

def _unique(xs):
    seen = set()
    for x in xs:
        if x not in seen:
            seen.add(x); yield x


# ==========================
# 컨테이너/포맷 판별 & 암호 감지
# ==========================
OOXML_CT_HINTS = (
    "application/vnd.ms-office.encryptedpackage",
    "application/vnd.openxmlformats-officedocument.encryptedpackage",
    "encryptedpackage",
)

def is_ooxml_zip(data: bytes) -> bool:
    return data.startswith(b'PK\x03\x04')

def is_pdf(data: bytes) -> bool:
    return data.startswith(b'%PDF-')

def is_ole_compound(data: bytes) -> bool:
    # D0 CF 11 E0 A1 B1 1A E1
    return data.startswith(bytes.fromhex('D0CF11E0A1B11AE1'))

OOXML_CT_HINTS = (
    "application/vnd.ms-office.encryptedpackage",
    "application/vnd.openxmlformats-officedocument.encryptedpackage",
    "encryptedpackage",
)

def is_ooxml_encrypted(data: bytes) -> bool:
    """
    OOXML(.xlsx/.docx) 암호 감지 (강화):
    - ZIP 엔트리에 EncryptionInfo / EncryptedPackage 존재
    - [Content_Types].xml에 encryptedPackage 타입 존재
    - (보강) ZIP 엔트리 구성이 '암호 파일 전형' (2~3개 엔트리) 인 경우
    """
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            names_list = [n for n in z.namelist()]
            names = {n.lower() for n in names_list}

            # 1) 전형적 암호 엔트리
            if "encryptioninfo" in names or "encryptedpackage" in names:
                return True

            # 2) Content_Types에 암호 패키지 타입
            try:
                ct = z.read("[Content_Types].xml").decode("utf-8", "ignore").lower()
                if any(h in ct for h in OOXML_CT_HINTS):
                    return True
            except KeyError:
                pass

            # 3) (보강 휴리스틱) 암호 OOXML의 전형적인 ZIP 레이아웃:
            #    [Content_Types].xml + EncryptionInfo + EncryptedPackage (대략 2~3개)
            #    즉, xl/workbook.xml, docProps/* 등 '정상 파트'가 전혀 없음
            normal_markers = (
                "xl/workbook.xml", "xl/worksheets/",
                "word/document.xml", "ppt/presentation.xml",
                "_rels/.rels", "docprops/core.xml"
            )
            has_normal = any(any(m in n.lower() for m in normal_markers) for n in names_list)
            if not has_normal and 1 <= len(names_list) <= 5:
                # 엔트리가 너무 적고(normal 파트가 없고) 암호 흔적이 거의 확실한 경우
                return True

            return False
    except zipfile.BadZipFile:
        return False

def is_ole_encrypted(data: bytes) -> bool:
    if not msoffcrypto:
        return False
    try:
        of = msoffcrypto.OfficeFile(io.BytesIO(data))
        return bool(getattr(of, "is_encrypted", False))
    except Exception:
        return False


# ==========================
# 텍스트 스캔 (요약 + 선택적 해시)
# ==========================
def scan_text(text: str, include_hashes: bool = False) -> Tuple[Dict[str,int], Dict[str, List[str]]]:
    counts = {
        "phone":0, "email":0,
        "rrn_candidate":0, "rrn_valid":0,
        "bizreg_candidate":0, "bizreg_valid":0,
        "card_candidate":0, "card_luhn":0,
        "account":0, "ip":0, "zip":0, "dob":0,
        "secrets_prefix":0, "secrets_entropy":0,
    }
    hashes: Dict[str, List[str]] = {} if include_hashes else {}

    if not text:
        return counts, hashes

    emails = list(_unique(REGEX["email"].findall(text)))
    phones = list(_unique(REGEX["phone"].findall(text)))
    ipv4s  = list(_unique(REGEX["ipv4"].findall(text)))
    ipv6s  = list(_unique(REGEX["ipv6"].findall(text)))
    zips   = list(_unique(REGEX["zip5"].findall(text)))
    dobs   = REGEX["dob"].findall(text)

    counts["email"] += len(emails)
    counts["phone"] += len(phones)
    counts["ip"]    += len(ipv4s) + len(ipv6s)
    counts["zip"]   += len(zips)
    counts["dob"]   += len(dobs)

    if include_hashes:
        hashes["email"] = [sha256_text(v) for v in emails]
        hashes["phone"] = [sha256_text(v) for v in phones]
        hashes["ip"]    = [sha256_text(v) for v in ipv4s + ipv6s]
        hashes["zip"]   = [sha256_text(v) for v in zips]

    rrns = list(_unique(REGEX["rrn"].findall(text)))
    counts["rrn_candidate"] += len(rrns)
    valid_rrn = [r for r in rrns if rrn_valid(r)]
    counts["rrn_valid"] += len(valid_rrn)
    if include_hashes:
        hashes["rrn_valid"] = [sha256_text(v) for v in valid_rrn]

    bizs = list(_unique(REGEX["bizreg"].findall(text)))
    counts["bizreg_candidate"] += len(bizs)
    valid_biz = [b for b in bizs if bizreg_valid(b)]
    counts["bizreg_valid"] += len(valid_biz)
    if include_hashes:
        hashes["bizreg_valid"] = [sha256_text(v) for v in valid_biz]

    for _ in REGEX["account"].finditer(text):
        counts["account"] += 1

    cards = list(_unique(REGEX["card_raw"].findall(text)))
    counts["card_candidate"] += len(cards)
    luhn_ok = [c for c in cards if luhn_check(c)]
    counts["card_luhn"] += len(luhn_ok)
    if include_hashes and luhn_ok:
        hashes["card_luhn"] = [sha256_text(v) for v in luhn_ok]

    secrets = list(_unique(REGEX["secret_prefix"].findall(text)))
    counts["secrets_prefix"] += len(secrets)
    he = [s for s in secrets if shannon_entropy(s) >= 3.5]
    counts["secrets_entropy"] += len(he)
    if include_hashes:
        hashes["secrets_prefix"]  = [sha256_text(v) for v in secrets]
        hashes["secrets_entropy"] = [sha256_text(v) for v in he]

    return counts, hashes


# ==========================
# PDF 암호 감지/추출기 (보강)
# ==========================
def _pdf_has_encrypt_marker(data: bytes) -> bool:
    """PDF 원문에 /Encrypt 키가 있으면 암호화로 간주(대부분 커버)."""
    if not data or not data.startswith(b"%PDF-"):
        return False
    return (b"/Encrypt" in data) or (b"/ENCRYPT" in data) or (b"/encrypt" in data)

def extract_text_from_pdf(data: bytes) -> str:
    # 0) 빠른 바이트 스캔
    try:
        if _pdf_has_encrypt_marker(data):
            raise PasswordProtectedError("pdf is password-protected")
    except PasswordProtectedError:
        raise
    except Exception:
        pass

    # 1) PyPDF2 경로(선택)
    if PdfReader is not None:
        try:
            r = PdfReader(io.BytesIO(data))
            if getattr(r, "is_encrypted", False):
                raise PasswordProtectedError("pdf is password-protected")
        except PasswordProtectedError:
            raise
        except Exception:
            pass

    # 2) pdfminer 텍스트 추출
    if not pdf_extract_text:
        return ""
    try:
        return pdf_extract_text(io.BytesIO(data)) or ""
    except Exception as e:
        msg = (str(e) or "").lower()
        if "password" in msg or "encrypted" in msg or "decrypt" in msg:
            raise PasswordProtectedError("pdf is password-protected")
        return ""


# ==========================
# 파일 추출기 (OOXML/OLE 암호 감지 포함)
# ==========================
MAX_FILE_BYTES = 25 * 1024 * 1024
FILE_EXT_WHITELIST = {
    "txt","csv","md","json","xml","html","log",
    "docx","xlsx","pdf",
    # 레거시 확장자도 허용(텍스트 추출은 기본판에서 미지원, 암호 판별만)
    "xls","doc"
}

def extract_text_from_docx(data: bytes) -> str:
    if not docx:
        return ""
    try:
        if is_ooxml_encrypted(data):
            raise PasswordProtectedError("docx is password-protected")
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text]
        for t in d.tables:
            for row in t.rows:
                cells = [c.text for c in row.cells if c.text]
                if cells: parts.append(" | ".join(cells))
        return "\n".join(parts)
    except PasswordProtectedError:
        raise
    except zipfile.BadZipFile:
        return ""
    except Exception as e:
        msg = (str(e) or "").lower()
        if "encrypted" in msg or "password" in msg or "decrypt" in msg:
            raise PasswordProtectedError("docx is password-protected")
        # 엄격 모드: OOXML인데 핵심 파트 없고 암호 흔적 있으면 암호로 승격
        if STRICT_OOXML_PASSWORD_HEURISTIC:
            try:
                z = zipfile.ZipFile(io.BytesIO(data))
                names = {n.lower() for n in z.namelist()}
                ct_has = False
                if "[content_types].xml" in names:
                    ct = z.read("[Content_Types].xml").decode("utf-8","ignore").lower()
                    ct_has = ("encryptedpackage" in ct)
                if ("word/document.xml" not in names) and ("encryptioninfo" in names or "encryptedpackage" in names or ct_has):
                    raise PasswordProtectedError("docx is password-protected")
            except PasswordProtectedError:
                raise
            except Exception:
                pass
        return ""

def extract_text_from_xlsx(data: bytes) -> str:
    if not load_workbook:
        return ""
    try:
        # A) 시그니처로 먼저 확인 (강화된 암호 감지)
        if is_ooxml_encrypted(data):
            raise PasswordProtectedError("xlsx is password-protected")

        # B) openpyxl 시도
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)

        # C) (보강) 정상 파트가 전혀 없는데 wb가 비정상적으로 비어있다면 암호 의심 승격
        try:
            if (not wb.worksheets) or (len(wb.worksheets) == 0):
                # zip 내부를 다시 확인해 암호 구조면 승격
                if is_ooxml_encrypted(data):
                    raise PasswordProtectedError("xlsx is password-protected")
        except Exception:
            pass

        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                vals = [str(v) for v in row if v is not None]
                if vals: parts.append(" , ".join(vals))
        return "\n".join(parts)

    except InvalidFileException as e:
        msg = (str(e) or "").lower()
        if "encrypted" in msg or "password" in msg or "decrypt" in msg:
            raise PasswordProtectedError("xlsx is password-protected")
        # D) openpyxl이 읽지 못했는데, OOXML 암호 구조면 강제 승격
        if is_ooxml_encrypted(data):
            raise PasswordProtectedError("xlsx is password-protected")
        return ""

    except zipfile.BadZipFile:
        return ""

    except KeyError:
        # E) 핵심 파트 없음 → 암호 구조면 승격
        if is_ooxml_encrypted(data):
            raise PasswordProtectedError("xlsx is password-protected")
        return ""

    except PasswordProtectedError:
        raise

    except Exception as e:
        msg = (str(e) or "").lower()
        if "encrypted" in msg or "password" in msg or "decrypt" in msg:
            raise PasswordProtectedError("xlsx is password-protected")

        # F) (엄격 모드) OOXML인데 정상 파트가 없고 암호 흔적이 있으면 승격
        if STRICT_OOXML_PASSWORD_HEURISTIC:
            try:
                with zipfile.ZipFile(io.BytesIO(data)) as z:
                    names_list = [n for n in z.namelist()]
                    names = {n.lower() for n in names_list}
                    ct_has = False
                    if "[content_types].xml" in names:
                        ct = z.read("[Content_Types].xml").decode("utf-8","ignore").lower()
                        ct_has = ("encryptedpackage" in ct)
                    if (
                        ("xl/workbook.xml" not in names) and
                        ("encryptioninfo" in names or "encryptedpackage" in names or ct_has)
                    ):
                        raise PasswordProtectedError("xlsx is password-protected")
            except PasswordProtectedError:
                raise
            except Exception:
                pass
        return ""

def extract_text_generic(name: str, data: bytes) -> str:
    if not data or len(data) > MAX_FILE_BYTES:
        return ""
    ext = (name.rsplit(".", 1)[-1].lower() if "." in name else "")

    try:
        # 컨테이너/매직 기반 빠른 분기
        if is_ooxml_zip(data):
            if ext == "docx":
                return extract_text_from_docx(data)
            if ext == "xlsx":
                return extract_text_from_xlsx(data)
            return ""

        if is_pdf(data):
            return extract_text_from_pdf(data)

        if is_ole_compound(data):
            if is_ole_encrypted(data):
                raise PasswordProtectedError("ole office file is password-protected")
            return ""  # (기본판) 레거시 텍스트 추출 미지원

        if ext in {"txt","csv","md","json","xml","html","log"}:
            return maybe_decode_text(data)

    except PasswordProtectedError:
        raise
    except Exception:
        return ""

    return ""


# ==========================
# 스캔 파이프라인
# ==========================
def handle_scan(parts: List[Dict[str, Any]], *, include_hashes: bool) -> Tuple[Dict[str,int], List[Dict[str,Any]], List[str], Dict[str,Any]]:
    summary: Dict[str, int] = {
        "phone":0, "email":0,
        "rrn_candidate":0, "rrn_valid":0,
        "bizreg_candidate":0, "bizreg_valid":0,
        "card_candidate":0, "card_luhn":0,
        "account":0, "ip":0, "zip":0, "dob":0,
        "secrets_prefix":0, "secrets_entropy":0,
    }
    files_meta: List[Dict[str, Any]] = []
    errors: List[str] = []
    details_total_hashes: Dict[str, List[str]] = {} if include_hashes else {}

    def _merge_counts(dst: Dict[str,int], src: Dict[str,int]):
        for k, v in src.items():
            dst[k] = int(dst.get(k,0)) + int(v)

    for p in parts:
        try:
            ptype = p.get("type")
            if ptype == "text":
                text = p.get("content","") or ""
                counts, hashes = scan_text(text, include_hashes=include_hashes)
                _merge_counts(summary, counts)
                if include_hashes:
                    for k, arr in hashes.items():
                        details_total_hashes.setdefault(k, []).extend(arr)

            elif ptype == "file":
                name = p.get("name") or "blob"
                b64  = p.get("bytes_base64") or ""
                data = base64.b64decode(b64) if b64 else b""

                meta = {
                    "name": name,
                    "size": len(data),
                    "sha256": sha256_bytes(data) if data else None,
                    "ext": (name.rsplit(".",1)[-1].lower() if "." in name else "")
                }

                try:
                    text = extract_text_generic(name, data)
                    counts, hashes = scan_text(text, include_hashes=include_hashes)
                    meta.update(counts)
                    if include_hashes:
                        meta["hashes"] = hashes
                except PasswordProtectedError as pe:
                    meta["error"] = "password_protected"
                    errors.append(f"file:{name} -> {type(pe).__name__}: {pe}")
                except Exception as e:
                    errors.append(f"file:{name} -> {type(e).__name__}: {e}")
                finally:
                    files_meta.append(meta)

            else:
                errors.append(f"unknown part type: {ptype}")
        except Exception as e:
            errors.append(f"{p.get('type','?')}:{p.get('name','?')} -> {type(e).__name__}: {e}")

    return summary, files_meta, errors, details_total_hashes


# ==========================
# 메인 루프
# ==========================
def main():
    while True:
        try:
            req = _read_msg()
        except Exception:
            sys.exit(0)

        kind = req.get("kind")

        if kind == "ping":
            _send_msg({"ok": True, "version": VERSION})
            continue

        if kind != "scan":
            _send_msg({"ok": False, "error": f"unknown kind: {kind}", "version": VERSION})
            continue

        options = req.get("options", {})
        include_hashes = bool(options.get("include_hashes", False))
        req_id = req.get("req_id")
        limits = req.get("limits", {})
        max_bytes = int(limits.get("max_bytes", MAX_FILE_BYTES))

        parts = req.get("parts", [])
        total_b64 = sum(len(p.get("bytes_base64","")) for p in parts if p.get("bytes_base64"))
        if total_b64 > max_bytes * 1.4:
            _send_msg({"ok": False, "error": "payload too large", "version": VERSION, "req_id": req_id})
            continue

        summary, files_meta, errors, total_hashes = handle_scan(parts, include_hashes=include_hashes)

        resp = {
            "ok": True,
            "version": VERSION,
            "req_id": req_id,
            "summary": summary,
            "files": files_meta,
            "errors": errors
        }
        if include_hashes:
            resp["details"] = {"hashes": total_hashes}

        _send_msg(resp)


if __name__ == "__main__":
    main()
